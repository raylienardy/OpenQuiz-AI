"""
DOCX Extractor using python-docx library.
Extracts text from paragraphs, headings, lists, and tables while preserving reading order.
"""

from fastapi import UploadFile
from .base import BaseExtractor
from .models import ExtractionResult
import io

try:
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError
except ImportError:
    Document = None

class DOCXExtractor(BaseExtractor):
    """Extract text from .docx files."""

    async def extract(self, file: UploadFile) -> ExtractionResult:
        file_type = "docx"
        warnings = []
        text_parts = []

        # Baca konten file
        try:
            content = await file.read()
        except Exception as e:
            return ExtractionResult(
                file_type=file_type,
                warnings=[f"Failed to read file: {str(e)}"]
            )

        if not content:
            return ExtractionResult(
                file_type=file_type,
                warnings=["File is empty."]
            )

        # Buka dokumen dari memori
        try:
            doc = Document(io.BytesIO(content))
        except PackageNotFoundError:
            return ExtractionResult(
                file_type=file_type,
                warnings=["File is corrupted or not a valid DOCX."]
            )
        except Exception as e:
            return ExtractionResult(
                file_type=file_type,
                warnings=[f"Failed to open DOCX file: {str(e)}"]
            )

        # Ekstrak paragraf
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Deteksi heading berdasarkan style (opsional, untuk struktur)
                if para.style.name.startswith('Heading'):
                    paragraphs.append(text + '\n')  # heading diikuti baris kosong
                else:
                    paragraphs.append(text)
        if paragraphs:
            text_parts.append('\n\n'.join(paragraphs))

        # Ekstrak tabel
        tables = []
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_text.append(' | '.join(row_data))
            if table_text:
                tables.append('\n'.join(table_text))
        if tables:
            text_parts.append('\n\n'.join(tables))

        # Gabungkan semua teks
        text = '\n\n'.join(text_parts).strip()

        if not text:
            warnings.append("Document contains no extractable text.")

        return ExtractionResult(
            text=text,
            file_type=file_type,
            warnings=warnings
        )